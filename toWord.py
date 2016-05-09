# coding=gbk

from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import datetime


def toWord(_conList_g, _conList_t):
    d = datetime.datetime.now()
    curTime = (d.strftime('%Y_%m_%d'))

    document = Document()

    style1 = document.styles['Body Text']
    font1 = style1.font
    font1.name = '����'
    font1.size = Pt(14)
    style2 = document.styles['Body Text 2']
    font2 = style2.font
    font2.name = '����'
    font2.size = Pt(12)
    style3 = document.styles['Body Text 3']
    font3 = style3.font
    font3.name = '����'
    font3.size = Pt(10.5)

    r = style1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '����')
    r = style2._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '����')
    r = style3._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '����')

    p = document.add_paragraph('')
    p.style = document.styles['Body Text']
    p.add_run("ȫְ��Ϣ").bold = True
    p.paragraph_format.line_spacing = 1

    for i in _conList_g:

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 2']
        p.add_run(i['Company']).bold = True
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("ѧ����")
        if(i['Schooling'] is not None):
            p.add_run(i['Schooling'])
        else:
            p.add_run("����")
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("רҵ��")
        if(i['Major'] is not None):
            p.add_run(i['Major'])
        else:
            p.add_run("����")
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("�����ص㣺")
        if(i['District'] is not None):
            p.add_run(i['District'])
        else:
            p.add_run("�������")
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("��Ƹ��λ��")
        if(i['RequireJob'] is not None):
            p.add_run(i['RequireJob'])
        else:
            p.add_run("�������")
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("����Ͷ�ݷ�ʽ��" + i['Method'])
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("��ֹ���ڣ�" + i['Deadline'][0:9])
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.paragraph_format.line_spacing = 1

    p = document.add_paragraph('')
    p.style = document.styles['Body Text']
    p.add_run("��ְ��Ϣ").bold = True
    p.paragraph_format.line_spacing = 1

    for i in _conList_t:

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 2']
        p.add_run(i['Company']).bold = True
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("ѧ����")
        if(i['Schooling'] is not None):
            p.add_run(i['Schooling'])
        else:
            p.add_run("����")
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("רҵ��")
        if(i['Major'] is not None):
            p.add_run(i['Major'])
        else:
            p.add_run("����")
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("�����ص㣺")
        if(i['District'] is not None):
            p.add_run(i['District'])
        else:
            p.add_run("�������")
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("��Ƹ��λ��")
        if(i['RequireJob'] is not None):
            p.add_run(i['RequireJob'])
        else:
            p.add_run("�������")
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("����Ͷ�ݷ�ʽ��" + i['Method'])
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.style = document.styles['Body Text 3']
        p.add_run("��ֹ���ڣ�" + i['Deadline'][0:9])
        p.paragraph_format.line_spacing = 1

        p = document.add_paragraph('')
        p.paragraph_format.line_spacing = 1

    document.save(curTime + '.docx')
